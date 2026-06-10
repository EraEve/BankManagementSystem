# -*- coding: utf-8 -*-
"""USER_MANUAL.md → 排版精美的 .docx 转换脚本"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SRC = os.path.expanduser("~/Desktop/BankManagementSystem_clone/USER_MANUAL.md")
DST = os.path.expanduser("~/Desktop/银行智能管理系统_使用说明文档.docx")

doc = Document()

# ── Page setup ────────────────────────────────────
for section in doc.sections:
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# ── Style helpers ─────────────────────────────────
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35

def set_font(run, name_cn, size, bold=False, color=None):
    """Set both Latin and East-Asian font on a run."""
    run.font.name = name_cn
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)
    if color:
        run.font.color.rgb = color

def add_title_page():
    """Cover / title page."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("银行智能管理系统")
    set_font(r, '黑体', 32, bold=True, color=RGBColor(8, 66, 152))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("使用说明文档  v2.2")
    set_font(r, '黑体', 20, bold=False, color=RGBColor(13, 110, 253))

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("数据结构与算法课程设计 · 第十组")
    set_font(r, '宋体', 12, color=RGBColor(100, 100, 100))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026-06-10  |  GitHub: EraEve/BankManagementSystem")
    set_font(r, '宋体', 10, color=RGBColor(130, 130, 130))

    doc.add_page_break()

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(8)
    # bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="8" w:space="4" w:color="084298"/>'
        f'</w:pBdr>')
    pPr.append(pBdr)
    r = p.add_run(text)
    set_font(r, '黑体', 18, bold=True, color=RGBColor(8, 66, 152))

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    set_font(r, '黑体', 14, bold=True, color=RGBColor(13, 110, 253))

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    set_font(r, '黑体', 12, bold=True, color=RGBColor(31, 41, 55))

def add_body(text):
    """Paragraph with mixed bold support."""
    if not text.strip():
        doc.add_paragraph()
        return
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.6)
    # Split by **bold** markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_font(r, '黑体', 10.5, bold=True)
        else:
            r = p.add_run(part)
            set_font(r, '宋体', 10.5)

def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    r = p.add_run("• " + text)
    set_font(r, '宋体', 10.5)

def add_code_block(lines):
    for l in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(1.0)
        # light blue-grey background via shading
        pPr = p._element.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4FA" w:val="clear"/>')
        pPr.append(shd)
        r = p.add_run(l[:110])
        set_font(r, '楷体', 9, color=RGBColor(51, 51, 51))

def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="C0C0C0"/>'
        f'</w:pBdr>')
    pPr.append(pBdr)

def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, '黑体', 9.5, bold=True, color=RGBColor(255, 255, 255))
        # dark blue background
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="084298" w:val="clear"/>')
        tcPr.append(shd)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            set_font(r, '宋体', 9)
            if ri % 2 == 0:
                tcPr = cell._element.get_or_add_tcPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F7FB" w:val="clear"/>')
                tcPr.append(shd)
    doc.add_paragraph()  # spacing after table

def add_blockquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    # left border via indent + grey text
    r = p.add_run(text)
    set_font(r, '楷体', 10, color=RGBColor(80, 80, 80))

# ── Parse & Build ─────────────────────────────────
with open(SRC, "r", encoding="utf-8") as f:
    lines = f.readlines()

add_title_page()

i = 0
in_toc = False
in_code = False
code_buf = []
in_table = False
table_buf = []
in_title_block = True
past_first_hr = False

while i < len(lines):
    line = lines[i].rstrip()

    # Skip title block until first ---
    if in_title_block:
        if line == "---":
            if past_first_hr:
                in_title_block = False
                i += 1
                continue
            past_first_hr = True
        i += 1
        continue

    # Skip TOC
    if line.startswith("## 目录"):
        in_toc = True; i += 1; continue
    if in_toc:
        if line.startswith("## "):
            in_toc = False
            # fall through to process header
        else:
            i += 1; continue

    # Code blocks
    if line.startswith("```"):
        if not in_code:
            in_code = True; code_buf = []; i += 1; continue
        else:
            in_code = False; add_code_block(code_buf); i += 1; continue
    if in_code:
        code_buf.append(line); i += 1; continue

    # Tables
    if "|" in line and line.strip().startswith("|"):
        if not in_table:
            in_table = True; table_buf = []
        table_buf.append(line); i += 1; continue
    if in_table:
        in_table = False
        headers = None; rows = []
        for r in table_buf:
            if re.match(r'^\|[\s\-:|]+\|$', r): continue
            cells = [c.strip() for c in r.split("|")[1:-1]]
            if headers is None:
                headers = cells
            else:
                rows.append(cells)
        if headers:
            add_table(headers, rows)
        table_buf = []; continue

    # Headers
    if line.startswith("### "):
        add_h3(line[4:].strip()); i += 1; continue
    if line.startswith("## "):
        txt = line[3:].strip()
        add_h2(txt); i += 1; continue
    if line.startswith("# "):
        txt = line[2:].strip()
        add_h1(txt); i += 1; continue

    # Empty
    if line.strip() == "":
        i += 1; continue

    # HR
    if line.strip() == "---":
        add_hr(); i += 1; continue

    # Bullets
    if line.strip().startswith("- ") or line.strip().startswith("* "):
        add_bullet(line.strip()[2:]); i += 1; continue

    # Numbered list
    if re.match(r'^\d+\.\s', line.strip()):
        add_bullet(line.strip()); i += 1; continue

    # Blockquote
    txt = line.strip()
    if txt.startswith("> "):
        add_blockquote(txt[2:]); i += 1; continue
    if txt.startswith(">"):
        add_blockquote(txt[1:]); i += 1; continue

    # Body
    if txt and not txt.startswith("<!--") and not txt.startswith("```"):
        add_body(txt)
    i += 1

# ── Clean up ──────────────────────────────────────
# Remove trailing empty paragraphs
while doc.paragraphs and doc.paragraphs[-1].text.strip() == "":
    p = doc.paragraphs[-1]._element
    p.getparent().remove(p)

doc.save(DST)
print(f"[OK] DOCX generated: {DST}")
print(f"  Paragraphs: {len(doc.paragraphs)}")
size_kb = os.path.getsize(DST) / 1024
print(f"  Size: {size_kb:.0f} KB")
